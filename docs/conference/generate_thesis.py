#!/usr/bin/env python3
"""Generate thesis.docx for 68th MIPT Conference."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# --- Page margins (2cm all sides) ---
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(2)
section.right_margin  = Cm(2)
section.top_margin    = Cm(2)
section.bottom_margin = Cm(2)

FONT_MAIN = "Times New Roman"
SIZE_MAIN = Pt(10)
SIZE_SMALL = Pt(9)

def para(text="", bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
         size=None, font=None, space_before=0, space_after=0):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.first_line_indent = Cm(0.75) if align == WD_ALIGN_PARAGRAPH.JUSTIFY and text else Cm(0)
    if text:
        run = p.add_run(text)
        run.bold   = bold
        run.italic = italic
        run.font.name = font or FONT_MAIN
        run.font.size = size or SIZE_MAIN
    return p

def para_mixed(parts, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=True, space_before=0, space_after=0):
    """parts: list of (text, bold, italic, size)"""
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.first_line_indent = Cm(0.75) if indent else Cm(0)
    for text, bold, italic, size in parts:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.name = FONT_MAIN
        run.font.size = size or SIZE_MAIN
    return p

def superscript_run(p, text, size=None):
    run = p.add_run(text)
    run.font.name = FONT_MAIN
    run.font.size = size or SIZE_MAIN
    run.font.superscript = True
    return run

# =====================================================================
# УДК
# =====================================================================
p_udk = doc.add_paragraph()
p_udk.alignment = WD_ALIGN_PARAGRAPH.LEFT
p_udk.paragraph_format.space_before = Pt(0)
p_udk.paragraph_format.space_after  = Pt(2)
r = p_udk.add_run("УДК 004.4'422")
r.font.name = FONT_MAIN
r.font.size = SIZE_MAIN

# =====================================================================
# Название
# =====================================================================
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.space_before = Pt(4)
p_title.paragraph_format.space_after  = Pt(4)
r = p_title.add_run(
    "Оптимизация последовательностей проходов компилятора\n"
    "с применением методов каузального вывода"
)
r.bold = True
r.font.name = FONT_MAIN
r.font.size = SIZE_MAIN

# =====================================================================
# Авторы
# =====================================================================
p_authors = doc.add_paragraph()
p_authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_authors.paragraph_format.space_before = Pt(2)
p_authors.paragraph_format.space_after  = Pt(0)

r1 = p_authors.add_run("Белонин В. В.")
r1.font.name = FONT_MAIN
r1.font.size = SIZE_MAIN
superscript_run(p_authors, "1", SIZE_SMALL)
p_authors.add_run(", ").font.name = FONT_MAIN
r2 = p_authors.add_run("Ефанов Н. Н.")
r2.font.name = FONT_MAIN
r2.font.size = SIZE_MAIN
superscript_run(p_authors, "1", SIZE_SMALL)

# =====================================================================
# Аффилиация
# =====================================================================
p_aff = doc.add_paragraph()
p_aff.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_aff.paragraph_format.space_before = Pt(0)
p_aff.paragraph_format.space_after  = Pt(6)
superscript_run(p_aff, "1", SIZE_SMALL)
r = p_aff.add_run(
    " Московский физико-технический институт (национальный исследовательский университет)"
)
r.font.name = FONT_MAIN
r.font.size = SIZE_SMALL

# =====================================================================
# ТЕКСТ
# =====================================================================

# Введение
para_mixed([
    ("Введение.", True, False, None),
    (" Проблема поиска порядка применения оптимизаций (", False, False, None),
    ("phase ordering problem", False, True, None),
    (") является одной из ключевых задач в оптимизирующих компиляторах: "
     "порядок применения оптимизационных проходов существенно влияет на качество "
     "генерируемого кода [1]. Стандартные решения — фиксированные конвейеры вида "
     "-O2, -Oz — выбраны эмпирически и не учитывают особенностей конкретной программы.", False, False, None),
])

para_mixed([
    ("", False, False, None),
    ("Распространённый подход к поиску улучшенных последовательностей — "
     "сбор статистики по случайным перестановкам проходов с последующим построением "
     "корреляционного графа переходов [1]. Ребро A→B в таком графе означает, что переход "
     "от прохода A к проходу B часто встречается в выигрышных последовательностях. "
     "Однако данный подход имеет принципиальный недостаток: корреляция не равна "
     "причинно-следственной связи.", False, False, None),
])

para_mixed([
    ("", False, False, None),
    ("Рассмотрим пример. Пусть проходы A и B оба полезны сами по себе (каждый уменьшает "
     "размер кода). Тогда любая последовательность, содержащая оба прохода, с высокой "
     "вероятностью окажется выигрышной — и в корреляционном графе появится ребро A→B. "
     "Однако порядок A→B может быть абсолютно эквивалентен порядку B→A: сам факт наличия "
     "обоих проходов объясняет успех, а не их взаимное расположение. Строя последовательности "
     "на основе корреляционного графа, мы рискуем воспроизводить ложные зависимости.", False, False, None),
])

para_mixed([
    ("", False, False, None),
    ("Именно здесь необходим ", False, False, None),
    ("каузальный вывод", False, True, None),
    (" (", False, False, None),
    ("causal inference", False, True, None),
    (") [2]: он позволяет оценить не «как часто A→B встречается в успешных последовательностях», "
     "а «улучшается ли результат именно из-за того, что A стоит перед B». "
     "Ключевое преимущество случайных перестановок проходов состоит в том, что они образуют "
     "приближённо рандомизированный эксперимент: каждый проход с одинаковой вероятностью "
     "оказывается до или после любого другого. Это является стандартным допущением для "
     "применения методов оценки среднего причинного эффекта (ATE) [2]. "
     "В данной работе предлагается применить методы каузального вывода для построения "
     "графа истинных причинно-следственных зависимостей между оптимизационными проходами "
     "компилятора с целью генерации оптимальных последовательностей для конкретных программ.", False, False, None),
])

# Метод
para_mixed([
    ("Метод.", True, False, None),
    (" Для пары проходов (A, B) вводим: группа «treatment» — прогоны, "
     "где A стоит непосредственно перед B; группа «control» — прогоны, где B стоит "
     "непосредственно перед A. Средний причинный эффект:", False, False, None),
])

# ATE formula
p_formula = doc.add_paragraph()
p_formula.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_formula.paragraph_format.first_line_indent = Cm(0)
p_formula.paragraph_format.space_before = Pt(2)
p_formula.paragraph_format.space_after  = Pt(2)
r = p_formula.add_run("ATE(A→B) = E[improvement | A перед B] − E[improvement | B перед A].")
r.font.name = FONT_MAIN
r.font.size = SIZE_MAIN
r.italic = True

para_mixed([
    ("", False, False, None),
    ("Ребро A→B включается в каузальный граф, если ATE(A→B) > 0 и скорректированное "
     "p-значение p", False, False, None),
    ("adj", False, True, None),
    (" < 0.05 (коррекция Бенджамини–Хохберга на множественное тестирование при 2 652 "
     "проверяемых парах [3]). Дополнительно вычисляется условный ATE (CATE) [4] — "
     "отдельно для каждого бенчмарка — для учёта гетерогенности эффектов. "
     "Применяется упрощённый PC-алгоритм [5] для проверки наличия опосредующих переменных. "
     "Также выполняется позиционный анализ: для каждого прохода сравниваются средние "
     "улучшения при расположении в начале и в конце последовательности.", False, False, None),
])

# Эксперимент
para_mixed([
    ("Эксперимент.", True, False, None),
    (" Экспериментальная выборка собрана на наборе ", False, False, None),
    ("cBench", False, True, None),
    (" [6] (19 программ из доменов телекоммуникаций, криптографии, мультимедиа и автомобильных "
     "вычислений). В качестве компилятора использован LLVM 16 с ", False, False, None),
    ("legacy pass manager", False, True, None),
    (". Рассматривался набор из 52 оптимизационных проходов (объединение стандартного "
     "конвейера -Oz для legacy PM и расширенного набора проходов). Для каждого из "
     "19 бенчмарков выполнено 5 000 прогонов со случайными подпоследовательностями "
     "переменной длины (от 10 до 30 проходов, сэмплирование без возвращения), итого "
     "95 000 прогонов. Использование подпоследовательностей переменной длины — в отличие "
     "от полных перестановок — обеспечивает необходимую вариацию присутствия/отсутствия "
     "каждого прохода, что является условием идентификации каузальных эффектов. "
     "Метрика: размер LLVM bitcode файла после оптимизации; базовая линия — флаг -Oz.", False, False, None),
])

# Результаты
para_mixed([
    ("Результаты.", True, False, None),
    (" Из 2 652 возможных пар проходов корреляционный граф содержит 2 652 ребра "
     "(100% пар встречались в выигрышных последовательностях). После применения "
     "каузального анализа с коррекцией FDR выявлено ", False, False, None),
    ("358 значимых причинно-следственных связей", True, False, None),
    (" (13.5% от всех пар): подавляющее большинство корреляций (86.5%) являются ложными "
     "и не несут информации об оптимальном порядке проходов.", False, False, None),
])

para_mixed([
    ("", False, False, None),
    ("Среди 19 исследованных программ наибольший потенциал оптимизации "
     "продемонстрировали (доля случайных последовательностей, превзошедших -Oz): "
     "dijkstra — 87.3%, sha — 80.9%, crc32 — 79.7%, bitcount — 77.6%, "
     "stringsearch — 73.0%. Позиционный анализ выявил 15 проходов "
     "со статистически значимыми предпочтениями: проходы ", False, False, None),
    ("strip, gvn, simplifycfg, instcombine", False, True, None),
    (" дают лучший результат в конце последовательности; "
     "проходы ", False, False, None),
    ("instnamer, loop-reduce, loop-unroll, loop-rotate", False, True, None),
    (" — в начале.", False, False, None),
])

# Заключение
para_mixed([
    ("Заключение.", True, False, None),
    (" Предложен метод на основе каузального вывода для построения графа истинных "
     "зависимостей между оптимизационными проходами компилятора LLVM. "
     "Из 2 652 корреляций отфильтровано 358 значимых причинно-следственных связей, "
     "которые используются для генерации оптимизированных последовательностей проходов. "
     "Разработаны стратегии обхода каузального графа: жадный алгоритм, поиск с лучом, "
     "MCTS. Применение CATE позволяет учитывать гетерогенные эффекты для разных "
     "классов программ. Направления дальнейшей работы: расширение набора бенчмарков, "
     "переход к метрике времени выполнения, применение структурных уравнений (SEM) "
     "для более детального моделирования зависимостей.", False, False, None),
])

# =====================================================================
# Литература
# =====================================================================
p_lit = doc.add_paragraph()
p_lit.alignment = WD_ALIGN_PARAGRAPH.LEFT
p_lit.paragraph_format.space_before = Pt(6)
p_lit.paragraph_format.space_after  = Pt(2)
p_lit.paragraph_format.first_line_indent = Cm(0)
r = p_lit.add_run("Литература")
r.bold = True
r.font.name = FONT_MAIN
r.font.size = SIZE_MAIN

refs = [
    "1. Fursin G. et al. Milepost GCC: Machine learning enabled self-tuning compiler // "
    "International Journal of Parallel Programming. 2011. Vol. 39(3). P. 296–327.",
    "2. Pearl J. Causality: Models, Reasoning and Inference. 2nd ed. "
    "Cambridge University Press, 2009. 464 p.",
    "3. Benjamini Y., Hochberg Y. Controlling the false discovery rate: a practical and "
    "powerful approach to multiple testing // Journal of the Royal Statistical Society. "
    "Series B. 1995. Vol. 57(1). P. 289–300.",
    "4. Athey S., Imbens G. W. Recursive partitioning for heterogeneous causal effects // "
    "Proceedings of the National Academy of Sciences. 2016. Vol. 113(27). P. 7353–7360.",
    "5. Spirtes P., Glymour C., Scheines R. Causation, Prediction, and Search. "
    "2nd ed. MIT Press, 2000. 543 p.",
    "6. Fursin G., Miranda C., Temam O. et al. A Practical Method For Quickly Evaluating Program "
    "Optimizations // HiPEAC'07. Ghent, 2007. P. 29–46.",
]
for ref in refs:
    p_ref = doc.add_paragraph()
    p_ref.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_ref.paragraph_format.space_before = Pt(0)
    p_ref.paragraph_format.space_after  = Pt(1)
    p_ref.paragraph_format.first_line_indent = Cm(0)
    r = p_ref.add_run(ref)
    r.font.name = FONT_MAIN
    r.font.size = SIZE_SMALL

doc.save("thesis.docx")
print("thesis.docx saved!")
